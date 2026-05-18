"""Tests for tools.utils — markdown→docx export and helpers."""
import pytest


def test_make_slug_basic():
    from tools.utils import make_slug
    assert make_slug("emotional dysregulation in ADHD") == "emotional-dysregulation-in-adhd"


def test_make_slug_strips_punctuation():
    from tools.utils import make_slug
    assert make_slug("What is depression?") == "what-is-depression"


def test_export_to_docx_renders_headings(monkeypatch, tmp_path):
    """H1/H2/H3 headings produce Word headings at matching levels."""
    md = (
        "# Top-level\n"
        "## Second\n"
        "### Third\n"
    )
    _patch_io(monkeypatch, tmp_path, md)
    from tools.utils import export_to_docx
    out_path = export_to_docx("dummy", "in.md", "out.docx")

    from docx import Document
    doc = Document(str(out_path))
    headings = [(p.text, p.style.name) for p in doc.paragraphs if p.style.name.startswith("Heading")]
    assert headings == [("Top-level", "Heading 1"), ("Second", "Heading 2"), ("Third", "Heading 3")]


def test_export_to_docx_renders_table(monkeypatch, tmp_path):
    """A GFM pipe table becomes a Word table with bold header row."""
    md = (
        "| Col A | Col B |\n"
        "| --- | --- |\n"
        "| x | y |\n"
        "| p | q |\n"
    )
    _patch_io(monkeypatch, tmp_path, md)
    from tools.utils import export_to_docx
    out_path = export_to_docx("dummy", "in.md", "out.docx")

    from docx import Document
    doc = Document(str(out_path))
    assert len(doc.tables) == 1
    tbl = doc.tables[0]
    assert tbl.rows[0].cells[0].text == "Col A"
    assert tbl.rows[0].cells[1].text == "Col B"
    assert tbl.rows[1].cells[0].text == "x"
    assert tbl.rows[2].cells[1].text == "q"
    # Header cell text run must be bold. `cell.text = ""` leaves an empty leading
    # run, then the bold run is appended — so check the run carrying "Col A".
    header_runs = tbl.rows[0].cells[0].paragraphs[0].runs
    assert any(r.text == "Col A" and r.bold for r in header_runs)


def test_export_to_docx_separator_does_not_eat_empty_data_row(monkeypatch, tmp_path):
    """F9 regression: `| | |` is a data row, not a separator, so it must survive."""
    md = (
        "| A | B |\n"
        "| --- | --- |\n"
        "| x | y |\n"
        "| | |\n"
        "| p | q |\n"
    )
    _patch_io(monkeypatch, tmp_path, md)
    from tools.utils import export_to_docx
    out_path = export_to_docx("dummy", "in.md", "out.docx")

    from docx import Document
    doc = Document(str(out_path))
    tbl = doc.tables[0]
    # Header + 3 data rows (the empty-cells row must be preserved).
    assert len(tbl.rows) == 4
    assert tbl.rows[2].cells[0].text == ""
    assert tbl.rows[2].cells[1].text == ""
    assert tbl.rows[3].cells[0].text == "p"


def test_export_to_docx_renders_bold(monkeypatch, tmp_path):
    """**bold** spans become bold runs inside a paragraph."""
    md = "This is **important** text.\n"
    _patch_io(monkeypatch, tmp_path, md)
    from tools.utils import export_to_docx
    out_path = export_to_docx("dummy", "in.md", "out.docx")

    from docx import Document
    doc = Document(str(out_path))
    para = doc.paragraphs[0]
    bold_runs = [r for r in para.runs if r.bold]
    assert any(r.text == "important" for r in bold_runs)


def test_export_to_docx_renders_bullet_list(monkeypatch, tmp_path):
    md = (
        "- first\n"
        "- second\n"
        "* third\n"
    )
    _patch_io(monkeypatch, tmp_path, md)
    from tools.utils import export_to_docx
    out_path = export_to_docx("dummy", "in.md", "out.docx")

    from docx import Document
    doc = Document(str(out_path))
    list_paragraphs = [p.text for p in doc.paragraphs if p.style.name == "List Bullet"]
    assert list_paragraphs == ["first", "second", "third"]


def test_export_to_docx_skips_horizontal_rule(monkeypatch, tmp_path):
    md = (
        "Before\n"
        "\n"
        "---\n"
        "\n"
        "After\n"
    )
    _patch_io(monkeypatch, tmp_path, md)
    from tools.utils import export_to_docx
    out_path = export_to_docx("dummy", "in.md", "out.docx")

    from docx import Document
    doc = Document(str(out_path))
    texts = [p.text for p in doc.paragraphs]
    assert "---" not in texts
    assert "Before" in texts
    assert "After" in texts


def _patch_io(monkeypatch, tmp_path, md: str) -> None:
    """Redirect read_output / get_output_dir to tmp_path with given markdown."""
    monkeypatch.setattr("tools.utils.read_output", lambda slug, name: md)
    monkeypatch.setattr("tools.utils.get_output_dir", lambda slug: tmp_path)
