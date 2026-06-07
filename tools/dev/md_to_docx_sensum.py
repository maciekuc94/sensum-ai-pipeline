# -*- coding: utf-8 -*-
"""
md_to_docx_sensum.py — konwerter dowolnego Markdown -> Word .docx w STYLU SENSUM.

JEDNO ZRODLO PRAWDY o brandowym docx. Styl marki jest zaszyty tutaj — nie odtwarzaj
go recznie i nie uzywaj tools/utils.export_to_docx (to nie-brandowy eksporter pipeline'u,
motyw Office, przykuty do outputs/videos_pl/).

DESIGN (2026-06-07 redesign — "editorial scientific journal, cieply"):
  Paleta: dwa kolory marki + ich odcienie UZYTE WYLACZNIE WEWNATRZ DOKUMENTU
  (nie w generowanych grafikach):
    - tlo  #F4E5CA  bezowy (page background)
    - atrament #582F0E  ciemny braz (tekst, linie, paski tabel)
    - #EBD9B8  glebszy bez  (chipy kodu, zebra w tabelach, tlo cytatu)  [odcien bezu]
    - #CDB488  przygaszony tan  (cienkie linie wewnatrz tabel)          [odcien bezu/brazu]
  Font: EB Garamond (tekst + naglowki), Consolas (kod).
  Hierarchia: masthead na str. 1 (logo + tytul + linia), H2 z brazowa linia,
  H3 z rozstrzelona kapitala, justowanie + automatyczne dzielenie wyrazow,
  chipy dla `code`, tabele z proporcjonalnymi kolumnami + zebra + powtarzanym
  naglowkiem (szerokie >=8 kol. ida w poziomie), cytat z lewym paskiem,
  stopka "SENSUM · <nr strony>", logo zmniejszane (z ~1.4 MB do ~100 KB).

Obsluga md: H1/H2/H3, **bold**, *italic*, `code`, listy (- / *), cytaty (>),
tabele GFM, <details>/<summary>, --- (linia ozdobna).

Uzycie (z korzenia repo):
  PYTHONIOENCODING=utf-8 python tools/dev/md_to_docx_sensum.py <plik.md> [<plik2.md> ...]
  ... [--out <katalog>]   # domyslnie .docx laduje obok zrodla
  ... [--no-logo]         # bez logo SENSUM na gorze
Wspiera tez glob: "docs/research/*.md"
"""
import argparse
import glob as globmod
import io
import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.enum.section import WD_SECTION, WD_ORIENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

ROOT = Path(__file__).resolve().parents[2]
LOGO = ROOT / "outputs" / "channel_assets" / "SENSUM_LOGO.png"

# --- Jedno zrodlo prawdy: paleta + fonty SENSUM ---
BROWN = RGBColor(0x58, 0x2F, 0x0E)
BEIGE = RGBColor(0xF4, 0xE5, 0xCA)
BROWN_HEX = "582F0E"
BEIGE_HEX = "F4E5CA"
TINT_HEX = "EBD9B8"   # glebszy bez: chipy kodu, zebra, tlo cytatu
HAIR_HEX = "CDB488"   # przygaszony tan: cienkie linie wewnatrz tabel
BODY_FONT = "EB Garamond"
MONO_FONT = "Consolas"

SEP_RE = re.compile(r"^\s*\|?(\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")
HR_RE = re.compile(r"^\s*([-*_])\1{2,}\s*$")
INLINE_RE = re.compile(r"(\*\*[^*]+\*\*|\*[^*]+\*|`[^`]+`)")

# kolejnosc krawedzi w pBdr narzucona przez schemat OOXML
_PBDR_ORDER = ["top", "left", "bottom", "right", "between", "bar"]


# ---------------------------------------------------------------- run / inline
def style_run(r, color=BROWN, font=BODY_FONT, bold=None, italic=None, size=None):
    r.font.color.rgb = color
    r.font.name = font
    if size is not None:
        r.font.size = Pt(size)
    rpr = r._element.get_or_add_rPr()
    rfonts = rpr.find(qn("w:rFonts"))
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.append(rfonts)
    rfonts.set(qn("w:ascii"), font)
    rfonts.set(qn("w:hAnsi"), font)
    rfonts.set(qn("w:cs"), font)
    if bold is not None:
        r.bold = bold
    if italic is not None:
        r.italic = italic


def run_shade(run, hexcolor):
    """Tlo pod runem (chip dla kodu)."""
    rpr = run._element.get_or_add_rPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hexcolor)
    rpr.append(shd)


def run_letterspace(run, twips):
    rpr = run._element.get_or_add_rPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:val"), str(twips))
    rpr.append(sp)


def shade_paragraph(p, hexcolor):
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    pPr.append(shd)


def para_border(p, edges):
    """edges: {'bottom': (hex, sz_eighths_pt, space_pt), 'left': (...)}"""
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.find(qn("w:pBdr"))
    if pbdr is None:
        pbdr = OxmlElement("w:pBdr")
        pPr.append(pbdr)
    for edge in _PBDR_ORDER:
        if edge not in edges:
            continue
        hexc, sz, space = edges[edge]
        el = OxmlElement(f"w:{edge}")
        el.set(qn("w:val"), "single")
        el.set(qn("w:sz"), str(sz))
        el.set(qn("w:space"), str(space))
        el.set(qn("w:color"), hexc)
        pbdr.append(el)


def add_inline(p, text, color=BROWN, italic=None, code_size=10):
    """Renderuje **bold**, *italic*, `code` (chip) w obrebie akapitu."""
    for part in INLINE_RE.split(text):
        if not part:
            continue
        if part.startswith("**") and part.endswith("**"):
            style_run(p.add_run(part[2:-2]), color, BODY_FONT, bold=True, italic=italic)
        elif part.startswith("`") and part.endswith("`"):
            r = p.add_run(" " + part[1:-1] + " ")  # spacje = padding chipa
            style_run(r, BROWN, MONO_FONT, italic=italic, size=code_size)
            run_shade(r, TINT_HEX)
        elif part.startswith("*") and part.endswith("*") and len(part) > 2:
            style_run(p.add_run(part[1:-1]), color, BODY_FONT, italic=True)
        else:
            style_run(p.add_run(part), color, BODY_FONT, italic=italic)


# ---------------------------------------------------------------- document setup
def set_page_background(doc, hexcolor):
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hexcolor)
    doc.element.insert(0, bg)
    doc.settings.element.append(OxmlElement("w:displayBackgroundShape"))


def enable_hyphenation(doc):
    s = doc.settings.element
    el = OxmlElement("w:autoHyphenation")
    el.set(qn("w:val"), "true")
    s.append(el)
    cl = OxmlElement("w:consecutiveHyphenLimit")
    cl.set(qn("w:val"), "2")
    s.append(cl)
    hz = OxmlElement("w:hyphenationZone")
    hz.set(qn("w:val"), "357")  # ~0.25 cala
    s.append(hz)


def setup_styles(doc):
    for section in doc.sections:
        section.top_margin = Inches(1.05)
        section.bottom_margin = Inches(0.95)
        section.left_margin = Inches(1.15)
        section.right_margin = Inches(1.15)

    normal = doc.styles["Normal"]
    normal.font.name = BODY_FONT
    normal.font.size = Pt(11)
    normal.font.color.rgb = BROWN
    pf = normal.paragraph_format
    pf.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    pf.line_spacing = 1.32
    pf.space_after = Pt(7)
    pf.space_before = Pt(0)
    pf.widow_control = True

    heads = {1: (24, 18, 6), 2: (15, 16, 4), 3: (12, 11, 3)}
    for lvl, (size, before, after) in heads.items():
        st = doc.styles[f"Heading {lvl}"]
        st.font.name = BODY_FONT
        st.font.bold = True
        st.font.color.rgb = BROWN
        st.font.size = Pt(size)
        spf = st.paragraph_format
        spf.space_before = Pt(before)
        spf.space_after = Pt(after)
        spf.line_spacing = 1.12
        spf.keep_with_next = True
        spf.widow_control = True


def add_footer(doc):
    p = doc.sections[0].footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    style_run(p.add_run("SENSUM"), BROWN, BODY_FONT, size=8.5)
    style_run(p.add_run("  ·  "), BROWN, BODY_FONT, size=8.5)
    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), " PAGE ")
    r = OxmlElement("w:r")
    rpr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    rf.set(qn("w:ascii"), BODY_FONT)
    rf.set(qn("w:hAnsi"), BODY_FONT)
    rpr.append(rf)
    col = OxmlElement("w:color")
    col.set(qn("w:val"), BROWN_HEX)
    rpr.append(col)
    sz = OxmlElement("w:sz")
    sz.set(qn("w:val"), "17")  # 8.5pt
    rpr.append(sz)
    r.append(rpr)
    t = OxmlElement("w:t")
    t.text = "1"
    r.append(t)
    fld.append(r)
    p._p.append(fld)


def add_logo(doc):
    if not LOGO.exists():
        return
    try:
        from PIL import Image
        img = Image.open(LOGO)
        if img.mode not in ("RGBA", "RGB"):
            img = img.convert("RGBA")
        max_side = 460
        if max(img.size) > max_side:
            ratio = max_side / max(img.size)
            img = img.resize((round(img.width * ratio), round(img.height * ratio)), Image.LANCZOS)
        buf = io.BytesIO()
        img.save(buf, format="PNG", optimize=True)
        buf.seek(0)
        src = buf
    except Exception:
        src = str(LOGO)
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(2)
    p.add_run().add_picture(src, width=Inches(1.15))


def add_masthead_title(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(7)
    p.paragraph_format.keep_with_next = True
    add_inline(p, text)
    for r in p.runs:
        r.font.size = Pt(23)
        r.bold = True
        r.font.color.rgb = BROWN
        run_letterspace(r, 4)
    para_border(p, {"bottom": (BROWN_HEX, 8, 6)})


# ---------------------------------------------------------------- tables
def set_cell_margins(table, top=60, bottom=60, left=110, right=110):
    mar = OxmlElement("w:tblCellMar")
    for edge, val in (("top", top), ("bottom", bottom), ("left", left), ("right", right)):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:w"), str(val))
        e.set(qn("w:type"), "dxa")
        mar.append(e)
    table._tbl.tblPr.append(mar)


def set_table_borders(table, outer_hex, inner_hex):
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        el = OxmlElement(f"w:{edge}")
        if edge == "insideV":
            el.set(qn("w:val"), "none")
        else:
            color, sz = (outer_hex, 8) if edge in ("top", "left", "bottom", "right") else (inner_hex, 4)
            el.set(qn("w:val"), "single")
            el.set(qn("w:sz"), str(sz))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), color)
        borders.append(el)
    table._tbl.tblPr.append(borders)


def shade(cell, hexcolor):
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:fill"), hexcolor)
    cell._tc.get_or_add_tcPr().append(shd)


def _cell_text(cell, txt, color=BROWN, bold=False, size=10.5):
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    cell.text = ""
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.line_spacing = 1.12
    add_inline(p, txt, color=color, code_size=size - 1)
    for r in p.runs:
        r.font.size = Pt(size)
        if bold:
            r.bold = True
        if color is not None:
            r.font.color.rgb = color


def _content_widths(padded, cols, avail_in):
    """Deterministyczne szerokosci kolumn: liczby waskie, tekst szeroki.
    Waga = przyciete maks. dlugosci tekstu w kolumnie (z naglowkiem), sciśnięta
    wykladnikiem 0.72, by jeden dlugi akapit nie zjadl calej szerokosci tabeli."""
    raw = [3.0] * cols
    for r in padded:
        for i in range(cols):
            txt = re.sub(r"[*`]", "", r[i]) if i < len(r) else ""
            raw[i] = max(raw[i], len(txt))
    weights = [min(x, 42) ** 0.72 for x in raw]
    total = sum(weights) or 1.0
    widths = [max(avail_in * w / total, 0.42) for w in weights]
    scale = avail_in / sum(widths)
    return [w * scale for w in widths]


def _apply_fixed_widths(table, widths_in):
    """Layout staly + jawne szerokosci (tblGrid + tcW + tblW), w poprawnej kolejnosci OOXML."""
    tblPr = table._tbl.tblPr
    grid = table._tbl.find(qn("w:tblGrid"))
    if grid is not None:
        for gc, w_in in zip(grid.findall(qn("w:gridCol")), widths_in):
            gc.set(qn("w:w"), str(int(round(w_in * 1440))))
    for i, w_in in enumerate(widths_in):
        for row in table.rows:
            row.cells[i].width = Inches(w_in)
    for el in tblPr.findall(qn("w:tblW")):
        tblPr.remove(el)
    tw = OxmlElement("w:tblW")
    tw.set(qn("w:type"), "dxa")
    tw.set(qn("w:w"), str(int(round(sum(widths_in) * 1440))))
    borders = tblPr.find(qn("w:tblBorders"))
    if borders is not None:
        borders.addprevious(tw)
    else:
        tblPr.append(tw)
    for el in tblPr.findall(qn("w:tblLayout")):
        tblPr.remove(el)
    tl = OxmlElement("w:tblLayout")
    tl.set(qn("w:type"), "fixed")
    if borders is not None:
        borders.addnext(tl)
    else:
        tblPr.append(tl)


def set_repeat_header(row):
    trPr = row._tr.get_or_add_trPr()
    th = OxmlElement("w:tblHeader")
    th.set(qn("w:val"), "true")
    trPr.append(th)
    ck = OxmlElement("w:cantSplit")
    ck.set(qn("w:val"), "true")
    trPr.append(ck)


def _no_split(row):
    trPr = row._tr.get_or_add_trPr()
    ck = OxmlElement("w:cantSplit")
    ck.set(qn("w:val"), "true")
    trPr.append(ck)


def add_table(doc, buf, font_size=10.5, avail_in=6.2, cell_mar=110):
    cols = max(len(r) for r in buf)
    padded = [r + [""] * (cols - len(r)) for r in buf]
    header, body = padded[0], padded[1:]
    table = doc.add_table(rows=1 + len(body), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_borders(table, BROWN_HEX, HAIR_HEX)
    set_cell_margins(table, left=cell_mar, right=cell_mar)

    for i, txt in enumerate(header):
        cell = table.rows[0].cells[i]
        shade(cell, BROWN_HEX)
        _cell_text(cell, txt, color=BEIGE, bold=True, size=font_size)
    for ri, row in enumerate(body, start=1):
        row_fill = TINT_HEX if (ri % 2 == 1) else BEIGE_HEX
        for i, txt in enumerate(row):
            cell = table.rows[ri].cells[i]
            shade(cell, row_fill)
            _cell_text(cell, txt, color=BROWN, bold=False, size=font_size)

    _apply_fixed_widths(table, _content_widths(padded, cols, avail_in))
    set_repeat_header(table.rows[0])
    for r in table.rows[1:]:
        _no_split(r)
    doc.add_paragraph().paragraph_format.space_after = Pt(4)


def _section_text_width_in(section):
    return (section.page_width - section.left_margin - section.right_margin) / 914400.0


def _start_landscape(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if sec.page_width < sec.page_height:
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.orientation = WD_ORIENT.LANDSCAPE
    sec.top_margin = Inches(0.8)
    sec.bottom_margin = Inches(0.8)
    sec.left_margin = Inches(0.8)
    sec.right_margin = Inches(0.8)


def _end_landscape(doc):
    sec = doc.add_section(WD_SECTION.NEW_PAGE)
    if sec.page_width > sec.page_height:
        sec.page_width, sec.page_height = sec.page_height, sec.page_width
    sec.orientation = WD_ORIENT.PORTRAIT
    sec.top_margin = Inches(1.05)
    sec.bottom_margin = Inches(0.95)
    sec.left_margin = Inches(1.15)
    sec.right_margin = Inches(1.15)


def _table_is_wide(tbl):
    return max(len(r) for r in tbl) >= 8


def _emit_table(doc, tbl):
    """Szeroka tabela (>=8 kolumn) -> sekcja pozioma, mniejsza czcionka; reszta -> pion."""
    if _table_is_wide(tbl):
        _start_landscape(doc)
        avail = _section_text_width_in(doc.sections[-1])
        add_table(doc, tbl, font_size=8.5, avail_in=avail, cell_mar=70)
        _end_landscape(doc)
    else:
        avail = _section_text_width_in(doc.sections[-1])
        add_table(doc, tbl, font_size=10.5, avail_in=avail, cell_mar=110)


def _landscape_item(blocks, j, n):
    """Czy blok j rozpoczyna element poziomy? Zwraca (typ, ile_blokow):
    'table' (sama szeroka tabela) lub 'caption_table' (podpis + szeroka tabela)."""
    kind, payload = blocks[j]
    if kind == "table" and _table_is_wide(payload):
        return ("table", 1)
    if (kind in ("p", "h2", "h3", "summary") and j + 1 < n
            and blocks[j + 1][0] == "table" and _table_is_wide(blocks[j + 1][1])):
        return ("caption_table", 2)
    return (None, 0)


# ---------------------------------------------------------------- parse + emit
def parse_blocks(md):
    """Markdown -> lista blokow [(kind, payload)] z mozliwoscia lookaheadu (podpis->tabela)."""
    blocks = []
    lines = md.splitlines()
    i, n = 0, len(lines)
    while i < n:
        s = lines[i].strip()
        if s.startswith("|") and s.endswith("|") and "|" in s[1:]:
            tbl = []
            while i < n:
                t = lines[i].strip()
                if not (t.startswith("|") and t.endswith("|") and "|" in t[1:]):
                    break
                if not SEP_RE.match(t):
                    tbl.append([c.strip() for c in t[1:-1].split("|")])
                i += 1
            if tbl:
                blocks.append(("table", tbl))
            continue
        i += 1
        if s == "" or s == ">":
            continue
        if s.startswith("<details") or s.startswith("</details"):
            continue
        m = re.match(r"<summary>(.*?)</summary>", s)
        if m:
            blocks.append(("summary", m.group(1)))
        elif HR_RE.match(s):
            blocks.append(("hr", None))
        elif s.startswith("### "):
            blocks.append(("h3", s[4:]))
        elif s.startswith("## "):
            blocks.append(("h2", s[3:]))
        elif s.startswith("# "):
            blocks.append(("h1", s[2:]))
        elif s.startswith("> "):
            blocks.append(("quote", s[2:]))
        elif s.startswith("- ") or s.startswith("* "):
            blocks.append(("bullet", s[2:]))
        else:
            blocks.append(("p", s))
    return blocks


def emit_simple(doc, kind, payload, state):
    if kind == "summary":
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        add_inline(p, "▸ " + payload)
        for r in p.runs:
            r.bold = True
    elif kind == "hr":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(10)
        para_border(p, {"bottom": (HAIR_HEX, 6, 1)})
    elif kind == "h3":
        p = doc.add_heading(level=3)
        add_inline(p, payload)
        for r in p.runs:
            run_letterspace(r, 14)
    elif kind == "h2":
        p = doc.add_heading(level=2)
        add_inline(p, payload)
        para_border(p, {"bottom": (BROWN_HEX, 8, 3)})
    elif kind == "h1":
        if not state["first_h1_done"]:
            add_masthead_title(doc, payload)
            state["first_h1_done"] = True
        else:
            p = doc.add_heading(level=1)
            add_inline(p, payload)
            para_border(p, {"bottom": (BROWN_HEX, 12, 4)})
    elif kind == "quote":
        p = doc.add_paragraph()
        pf = p.paragraph_format
        pf.left_indent = Inches(0.38)
        pf.right_indent = Inches(0.25)
        pf.space_before = Pt(4)
        pf.space_after = Pt(8)
        para_border(p, {"left": (BROWN_HEX, 18, 10)})
        shade_paragraph(p, TINT_HEX)
        add_inline(p, payload, italic=True)
    elif kind == "bullet":
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        pf = p.paragraph_format
        pf.left_indent = Inches(0.42)
        pf.first_line_indent = Inches(-0.22)
        pf.space_after = Pt(4)
        pf.line_spacing = 1.26
        style_run(p.add_run("•  "))
        add_inline(p, payload)
    else:  # "p"
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        add_inline(p, payload)


# ---------------------------------------------------------------- convert
def convert(src: Path, out: Path, logo: bool = True):
    md = src.read_text(encoding="utf-8")
    doc = Document()
    setup_styles(doc)
    set_page_background(doc, BEIGE_HEX)
    enable_hyphenation(doc)
    add_footer(doc)
    if logo:
        add_logo(doc)

    state = {"first_h1_done": False}
    blocks = parse_blocks(md)
    i, n = 0, len(blocks)
    while i < n:
        kind, payload = blocks[i]
        litype, _ = _landscape_item(blocks, i, n)
        if litype:
            # grupuj sasiadujace elementy poziome w JEDNEJ sekcji (bez pustych stron)
            _start_landscape(doc)
            while i < n:
                t, c = _landscape_item(blocks, i, n)
                if not t:
                    break
                avail = _section_text_width_in(doc.sections[-1])
                if t == "caption_table":
                    emit_simple(doc, blocks[i][0], blocks[i][1], state)
                    add_table(doc, blocks[i + 1][1], font_size=8.5, avail_in=avail, cell_mar=70)
                else:
                    add_table(doc, blocks[i][1], font_size=8.5, avail_in=avail, cell_mar=70)
                i += c
            if i < n:  # wroc do pionu tylko gdy jest dalsza tresc (inaczej pusta strona na koncu)
                _end_landscape(doc)
            continue
        if kind == "table":  # waska tabela -> pion
            _emit_table(doc, payload)
            i += 1
            continue
        emit_simple(doc, kind, payload, state)
        i += 1

    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out))
    return out


def _expand(paths):
    files = []
    for pat in paths:
        if any(ch in pat for ch in "*?["):
            files.extend(sorted(globmod.glob(pat)))
        else:
            files.append(pat)
    return [Path(f) for f in files]


def main():
    ap = argparse.ArgumentParser(description="Konwerter Markdown -> docx w stylu SENSUM.")
    ap.add_argument("paths", nargs="+", help="pliki .md lub glob (np. docs/research/*.md)")
    ap.add_argument("--out", default=None, help="katalog wyjsciowy (domyslnie obok zrodla)")
    ap.add_argument("--no-logo", action="store_true", help="bez logo SENSUM na gorze")
    args = ap.parse_args()

    srcs = [p for p in _expand(args.paths) if p.suffix.lower() == ".md" and p.exists()]
    if not srcs:
        print("Brak plikow .md do konwersji.")
        return
    out_dir = Path(args.out) if args.out else None
    for src in srcs:
        out = (out_dir / (src.stem + ".docx")) if out_dir else src.with_suffix(".docx")
        try:
            written = convert(src, out, logo=not args.no_logo)
            print(f"OK  {src.name} -> {written}  ({written.stat().st_size} B)")
        except Exception as ex:
            print(f"FAIL {src.name}: {ex}")
    print("DONE")


if __name__ == "__main__":
    main()
