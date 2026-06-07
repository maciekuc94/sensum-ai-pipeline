"""Shared utilities for all agent scripts in the WAT framework."""

import json
import os
import re
from datetime import datetime
from pathlib import Path

from dotenv import load_dotenv
from slugify import slugify


# Load environment variables from .env file at project root
# ---------------------------------------------------------------------------
# Shared visual brand constants (used by agent5_visuals and agent6_images)
# ---------------------------------------------------------------------------

STYLE_SUFFIX = (
    "minimalist high-contrast ink illustration on a flat solid sage beige background (#F4E5CA) — "
    "the background is one continuous solid sage beige color with no texture, no paper grain, no mottling, "
    "no aged-paper effects, no parchment look, no fibers, no stains, no discoloration, "
    "no border, no frame, no decorative outline around the image, no inner rectangle or panel containing the illustration, "
    "full-bleed composition where the background extends edge-to-edge, "
    "color palette strictly limited to #582F0E dark brown ink lines on #F4E5CA sage beige — no other colors whatsoever, "
    "technique: detailed cross-hatching for depth and shadow, fine-liner ink sketch, 2D perspective, heavy negative space, "
    "style: 19th-century scientific journal engraving, zero photorealism, no 3D effects, no gradients, no glows, no blurs, "
    "no green, no golden ochre, no moss green, no watercolor, no color fills, "
    "absolutely no text, no words, no letters, no numbers, no labels, no captions anywhere in the image, "
    "framing rule: if any human figure appears, the entire head must be fully visible inside the frame with generous headroom above — "
    "never crop the top of the head, never let the head touch or extend past the top edge of the frame, "
    "if the scene calls for a hand or close-up detail only, show ONLY the hand or detail with no head or body visible, "
    "16:9 aspect ratio"
)

CHARACTER_DESCRIPTION = (
    "a completely androgynous faceless gender-neutral human mannequin figure with a smooth featureless blank oval head — "
    "the head is a plain empty oval shape like an unpainted mannequin or blank egg, "
    "absolutely no facial features of any kind on the head, "
    "no eyes, no nose, no mouth, no ears, no hair, "
    "the body is fully androgynous: narrow symmetrical shoulders equal in width to the hips, slim flat chest with no breast or pectoral definition whatsoever, "
    "straight waist with no curve or hip flare, legs equal length, completely ambiguous gender with no masculine or feminine body markers, "
    "drawn entirely in fine-liner ink lines in deep espresso brown (#582F0E), "
    "cross-hatching used for depth and shading on the body, no color fills, "
    "19th-century scientific illustration style, 2D flat"
)


def _load_env():
    """Load .env file from project root."""
    env_path = Path(__file__).parent.parent / ".env"
    load_dotenv(env_path, override=False)


_load_env()


def get_env(key: str) -> str:
    """
    Get an environment variable with validation.

    Args:
        key: The environment variable name

    Returns:
        The environment variable value

    Raises:
        EnvironmentError: If the variable is missing or empty
    """
    value = os.getenv(key)
    if not value:
        raise EnvironmentError(
            f"Environment variable '{key}' is missing or empty. "
            f"Please set it in .env file at project root or in your shell."
        )
    return value


def make_slug(topic: str) -> str:
    """
    Convert a topic string to a filesystem-safe slug.

    Examples:
        "emotional dysregulation in ADHD" -> "emotional_dysregulation_in_adhd"
        "What is depression?" -> "what_is_depression"

    Args:
        topic: The topic string to slugify

    Returns:
        A filesystem-safe slug
    """
    return slugify(topic, separator="_", lowercase=True)


def next_output_number() -> int:
    """Return the next sequential video output number by scanning outputs/ for N_ prefixes."""
    outputs_dir = Path(__file__).parent.parent / "outputs" / "videos_pl"
    max_n = 0
    if outputs_dir.exists():
        for entry in outputs_dir.iterdir():
            if entry.is_dir():
                m = re.match(r"^(\d+)_", entry.name)
                if m:
                    n = int(m.group(1))
                    if n > max_n:
                        max_n = n
    return max_n + 1


def get_output_dir(slug: str) -> Path:
    """
    Get the output directory for a given slug, creating it if needed.

    Creates the slug directory and standard subdirectories on first access.

    Args:
        slug: The slug identifier for this task

    Returns:
        Path to the output directory (e.g., outputs/emotional-dysregulation-in-adhd)
    """
    output_dir = Path(__file__).parent.parent / "outputs" / "videos_pl" / slug
    # Scaffold the same skeleton as slug 1. No images_grain — grain is off by
    # default; agent6_images.py creates images_grain/ itself only when grain is
    # actually applied. The edit/ folder is created lazily by agent_align.py.
    for subdir in ("images", "md", "docx", "thumbnails", "thumbnails_grain", "thumbnails_no_grain", "voiceover"):
        (output_dir / subdir).mkdir(parents=True, exist_ok=True)
    return output_dir


def write_output(slug: str, filename: str, content: str) -> Path:
    """
    Write content to a file in the output directory.

    Args:
        slug: The slug identifier for this task
        filename: The filename to write (e.g., "summary.md", "images/thumb.png")
        content: The content to write (string)

    Returns:
        Path to the written file
    """
    output_dir = get_output_dir(slug)
    file_path = output_dir / filename

    # Create parent directories if needed
    file_path.parent.mkdir(parents=True, exist_ok=True)

    file_path.write_text(content, encoding="utf-8")
    return file_path


def read_output(slug: str, filename: str) -> str:
    """
    Read content from a file in the output directory.

    Args:
        slug: The slug identifier for this task
        filename: The filename to read (e.g., "summary.md")

    Returns:
        The file contents as a string

    Raises:
        FileNotFoundError: If the file doesn't exist with a helpful message
    """
    output_dir = get_output_dir(slug)
    file_path = output_dir / filename

    if not file_path.exists():
        raise FileNotFoundError(
            f"Output file not found: {file_path}\n"
            f"Expected file in outputs/{slug}/{filename}\n"
            f"Make sure the file was created by a previous step."
        )

    return file_path.read_text(encoding="utf-8")


# NOTE: query_claude() (direct Anthropic API) was removed on 2026-05-29. The
# pipeline no longer makes any Claude/Anthropic API call — all Claude work runs
# in-session in Claude Code on Opus 4.8 via slash commands (/draft, /hook,
# /visuals, /package). ANTHROPIC_API_KEY is no longer required. Text models
# that remain are Gemini (Vertex AI) for research and image rendering only.


def query_gemini_text(
    prompt: str,
    model: str,
    max_tokens: int,
    step_label: str = "",
) -> tuple[str, dict]:
    """Call Gemini via Vertex AI (google-genai SDK) and return (response_text, usage_dict).

    Plain text generation — no Google Search grounding.
    Retries up to 3 times on quota/server errors with exponential backoff at 15s, 30s, 60s.
    """
    import time
    import google.genai as genai
    from google.genai import types

    project = get_env("GOOGLE_CLOUD_PROJECT")
    location = "global"
    label = f" — {step_label}" if step_label else ""
    print(f"  Querying {model}{label}...")

    client = genai.Client(vertexai=True, project=project, location=location)

    for attempt in range(1, 5):
        try:
            response = client.models.generate_content(
                model=model,
                contents=prompt,
                config=types.GenerateContentConfig(
                    max_output_tokens=max_tokens,
                ),
            )
            usage = {"model": model, "input_tokens": 0, "output_tokens": 0}
            if hasattr(response, "usage_metadata") and response.usage_metadata:
                usage["input_tokens"] = getattr(response.usage_metadata, "prompt_token_count", 0) or 0
                usage["output_tokens"] = getattr(response.usage_metadata, "candidates_token_count", 0) or 0
            text = response.text
            if text is None:
                # Try to recover partial text from candidates (e.g. MAX_TOKENS truncation)
                try:
                    if response.candidates and response.candidates[0].content.parts:
                        text = "".join(
                            p.text for p in response.candidates[0].content.parts
                            if hasattr(p, "text") and p.text
                        ) or None
                except Exception:
                    pass
            if text is None:
                finish = getattr(response.candidates[0] if response.candidates else None, "finish_reason", "unknown") if hasattr(response, "candidates") else "unknown"
                raise ValueError(f"Gemini returned None text (finish_reason={finish}). Possible safety block or empty candidate.")
            return text, usage
        except Exception as exc:
            status = getattr(exc, "status_code", None) or getattr(exc, "code", None)
            is_retryable = status in (429, 503) or "quota" in str(exc).lower() or "unavailable" in str(exc).lower()
            if is_retryable and attempt < 4:
                wait = 15 * (2 ** (attempt - 1))
                print(f"  Gemini rate limited — waiting {wait}s before retry {attempt}/3...")
                time.sleep(wait)
                continue
            raise


def log_cost(slug: str, agent: str, data: dict) -> None:
    """No-op. cost_log.json is no longer written (not needed per user, 2026-05-29).

    Kept as an importable no-op so existing callers (agent_align) don't
    break. Remove the calls entirely if/when convenient.
    """
    return


def read_script_docx_text(path) -> str:
    """Extract plain narration body from a script .docx, skipping heading paragraphs."""
    from docx import Document
    doc = Document(Path(path))
    lines = []
    for para in doc.paragraphs:
        if para.style.name.startswith("Heading"):
            continue
        lines.append(para.text)
    return "\n".join(lines).strip()


def export_to_docx(
    slug: str,
    md_filename: str,
    docx_filename: str,
    sentence_per_line: bool = False,
    no_spacing: bool = False,
    preserve_blank_lines: bool = False,
) -> Path:
    """Convert a markdown file in the output dir to a .docx Word document.

    Handles: H1/H2/H3 headings, bullet lines (- ...), bold (**text**), GFM
    pipe tables, and plain paragraphs. Horizontal rules (---) are skipped.

    When ``sentence_per_line`` is True, plain paragraphs are split on sentence
    boundaries (.!?) so each sentence becomes its own docx paragraph. Headings,
    bullets, and tables are unaffected.

    Returns the path to the written .docx file.
    """
    _sentence_split_re = re.compile(r"(?<=[.!?])\s+(?=\S)")
    from docx import Document
    from docx.shared import Pt

    def _apply_spacing(para):
        if no_spacing:
            para.paragraph_format.space_before = Pt(0)
            para.paragraph_format.space_after = Pt(0)

    md_text = read_output(slug, md_filename)
    doc = Document()

    # Remove default empty paragraph Word adds
    for para in doc.paragraphs:
        p = para._element
        p.getparent().remove(p)

    def _add_inline(paragraph, text: str) -> None:
        """Add text to a paragraph, converting **bold** and <u>underline</u> spans."""
        parts = re.split(r"(\*\*[^*]+\*\*|<u>[^<]+</u>)", text)
        for part in parts:
            if part.startswith("**") and part.endswith("**"):
                run = paragraph.add_run(part[2:-2])
                run.bold = True
            elif part.startswith("<u>") and part.endswith("</u>"):
                run = paragraph.add_run(part[3:-4])
                run.underline = True
            elif part:
                paragraph.add_run(part)

    # GFM table separator: each cell must contain >=3 dashes (optionally with colons for alignment).
    # Prior pattern `^[\|\s:\-]+$` also swallowed empty data rows like `| | |`.
    _separator_re = re.compile(r"^\s*\|?(\s*:?-{3,}:?\s*\|)+\s*:?-{3,}:?\s*\|?\s*$")

    def _parse_row(line: str) -> list[str]:
        # Strip leading/trailing pipe, split on |, trim cells.
        inner = line.strip()
        if inner.startswith("|"):
            inner = inner[1:]
        if inner.endswith("|"):
            inner = inner[:-1]
        return [c.strip() for c in inner.split("|")]

    table_buffer: list[list[str]] = []

    def _flush_table() -> None:
        if not table_buffer:
            return
        col_count = max(len(r) for r in table_buffer)
        # Pad short rows so every row has col_count cells.
        padded = [r + [""] * (col_count - len(r)) for r in table_buffer]
        header = padded[0]
        body = padded[1:]
        try:
            table = doc.add_table(rows=1 + len(body), cols=col_count)
            table.style = "Light Grid Accent 1"
        except KeyError:
            # Style may not exist on minimal docx templates; fall back to default.
            table = doc.add_table(rows=1 + len(body), cols=col_count)
        # Header row — bold.
        for i, cell_text in enumerate(header):
            cell = table.rows[0].cells[i]
            cell.text = ""
            para = cell.paragraphs[0]
            run = para.add_run(cell_text)
            run.bold = True
        # Body rows.
        for r, row in enumerate(body, start=1):
            for i, cell_text in enumerate(row):
                cell = table.rows[r].cells[i]
                cell.text = ""
                para = cell.paragraphs[0]
                _add_inline(para, cell_text)
        _apply_spacing(doc.add_paragraph())  # spacer after table
        table_buffer.clear()

    for line in md_text.splitlines():
        stripped = line.strip()

        # Table line (must come before the --- check so separator rows aren't dropped).
        if stripped.startswith("|") and stripped.endswith("|") and "|" in stripped[1:]:
            if _separator_re.match(stripped):
                # GFM header separator — discard, don't add as a row.
                continue
            table_buffer.append(_parse_row(stripped))
            continue

        # Anything non-table flushes any in-progress table.
        _flush_table()

        if stripped.startswith("### "):
            _apply_spacing(doc.add_heading(stripped[4:], level=3))
        elif stripped.startswith("## "):
            _apply_spacing(doc.add_heading(stripped[3:], level=2))
        elif stripped.startswith("# "):
            _apply_spacing(doc.add_heading(stripped[2:], level=1))
        elif stripped.startswith("- ") or stripped.startswith("* "):
            para = doc.add_paragraph(style="List Bullet")
            _add_inline(para, stripped[2:])
            _apply_spacing(para)
        elif stripped.startswith("---"):
            # Horizontal rule — skip
            continue
        elif stripped == "":
            if preserve_blank_lines:
                blank = doc.add_paragraph()
                _apply_spacing(blank)
            else:
                continue
        else:
            if sentence_per_line:
                for sentence in _sentence_split_re.split(stripped):
                    sentence = sentence.strip()
                    if not sentence:
                        continue
                    para = doc.add_paragraph()
                    _add_inline(para, sentence)
                    _apply_spacing(para)
                    if preserve_blank_lines and re.search(r"\[visual pause\]", sentence, re.IGNORECASE):
                        extra = doc.add_paragraph()
                        _apply_spacing(extra)
            else:
                para = doc.add_paragraph()
                _add_inline(para, stripped)
                _apply_spacing(para)
                if preserve_blank_lines and re.search(r"\[visual pause\]", stripped, re.IGNORECASE):
                    extra = doc.add_paragraph()
                    _apply_spacing(extra)

    _flush_table()  # in case the file ends with a table

    output_dir = get_output_dir(slug)
    output_path = output_dir / docx_filename
    doc.save(str(output_path))
    return output_path


# ---------------------------------------------------------------------------
# Image post-processing (shared by agent6_images and agent7_thumbnails)
# ---------------------------------------------------------------------------

TARGET_BACKGROUND_RGB = (244, 229, 202)  # #F4E5CA sage beige
TARGET_IMAGE_SIZE = (1920, 1080)
BACKGROUND_THRESHOLD_DEFAULT = 170
# Threshold 170 catches off-brand beiges (Gemini drifts to e.g. #E7D7B5 / #F4DCB5,
# min channel ~175-181) that the old >240 threshold silently missed. Ink (#582F0E)
# has min channel ~14 and cross-hatch mid-tones land ~100-160, so 170 leaves
# shading intact. History note: a prior bug used 240 vs 170 between agent9 and
# agent7 — keep both agents on the same threshold by importing from here.


def resize_to_target(
    image_path,
    *,
    target_size: tuple[int, int] = TARGET_IMAGE_SIZE,
    background: tuple[int, int, int] = TARGET_BACKGROUND_RGB,
) -> None:
    """Scale image to fit inside target_size, pad remainder with background — no crop, no stretch."""
    from PIL import Image, ImageOps
    img = Image.open(str(image_path)).convert("RGB")
    if img.size != target_size:
        img = ImageOps.pad(img, target_size, color=background, centering=(0.5, 0.5))
        img.save(str(image_path))


def enforce_background_color(
    image_path,
    *,
    threshold: int = BACKGROUND_THRESHOLD_DEFAULT,
    background: tuple[int, int, int] = TARGET_BACKGROUND_RGB,
) -> None:
    """Replace all light background pixels with the exact brand color.

    Any pixel whose every channel exceeds `threshold` is treated as background
    and replaced with `background`. See module-level note on threshold=170.
    """
    import numpy as np
    from PIL import Image

    arr = np.array(Image.open(str(image_path)).convert("RGB"))
    mask = (arr[:, :, 0] > threshold) & (arr[:, :, 1] > threshold) & (arr[:, :, 2] > threshold)
    arr[mask] = background
    Image.fromarray(arr).save(str(image_path))


def make_transparent(image_path, *, threshold: int = BACKGROUND_THRESHOLD_DEFAULT) -> None:
    """Convert light background pixels to transparent (alpha=0). Keeps dark ink pixels opaque.

    Any pixel where all channels exceed `threshold` becomes fully transparent.
    Output is RGBA PNG. Threshold mirrors enforce_background_color for consistency.
    """
    import numpy as np
    from PIL import Image

    arr = np.array(Image.open(str(image_path)).convert("RGBA"))
    mask = (arr[:, :, 0] > threshold) & (arr[:, :, 1] > threshold) & (arr[:, :, 2] > threshold)
    arr[mask, 3] = 0
    arr[~mask, 3] = 255
    Image.fromarray(arr).save(str(image_path))


# ---------------------------------------------------------------------------
# Film grain (shared by agent6_images, agent7_thumbnails, tools/dev/add_grain)
# ---------------------------------------------------------------------------

GRAIN_INTENSITY_DEFAULT = 12  # SENSUM standard — Gaussian std-dev on 0–255 scale


def add_grain(
    image_path,
    *,
    intensity: int = GRAIN_INTENSITY_DEFAULT,
    out_path=None,
    rng_seed: int | None = None,
) -> Path:
    """
    Apply Gaussian film grain to a single PNG.

    When `out_path` is None the image is overwritten in place; otherwise the
    result is written to `out_path` (parent directory must already exist).
    Set `rng_seed` for deterministic noise — the batch tool uses 42 so re-runs
    are reproducible.
    """
    import numpy as np
    from PIL import Image

    src = Path(image_path)
    dst = Path(out_path) if out_path is not None else src
    img = Image.open(str(src))
    mode = img.mode if img.mode in ("RGB", "RGBA") else "RGB"
    arr = np.array(img.convert(mode), dtype=np.int16)
    rng = np.random.default_rng(rng_seed) if rng_seed is not None else np.random.default_rng()
    if mode == "RGBA":
        noise = rng.normal(0, intensity, arr[:, :, :3].shape).astype(np.int16)
        arr[:, :, :3] = np.clip(arr[:, :, :3] + noise, 0, 255)
    else:
        noise = rng.normal(0, intensity, arr.shape).astype(np.int16)
        arr = np.clip(arr + noise, 0, 255)
    Image.fromarray(arr.astype(np.uint8), mode).save(str(dst))
    return dst
